# app.py
from flask import Flask, request, render_template, send_file, jsonify
import os
import subprocess
import uuid
import time
import shutil
import logging

# Set up logging
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
                    handlers=[logging.FileHandler("app.log"), logging.StreamHandler()])
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload size

# Create necessary directories if they don't exist
for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
    if not os.path.exists(folder):
        os.makedirs(folder)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and file.filename.lower().endswith(('.docx', '.doc')):
        # Generate unique filename to prevent collisions
        unique_id = str(uuid.uuid4())
        timestamp = int(time.time())
        file_id = f"{timestamp}_{unique_id}"
        
        # Save uploaded file with the unique ID
        input_filename = f"{file_id}_{file.filename}"
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], input_filename)
        file.save(input_path)
        
        # Create output filename and path
        output_filename = f"{file_id}_{os.path.splitext(file.filename)[0]}.pdf"
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        
        try:
            # Try conversion methods
            logger.info(f"Starting conversion for {file.filename}")
            conversion_success = False
            error_messages = []
            
            # Try with PDF converters in order of preference
            conversion_methods = [
                ('pandoc', convert_with_pandoc),
                ('libreoffice', convert_with_libreoffice),
                ('unoconv', convert_with_unoconv),
                ('python-docx-pypdf', convert_with_python_docx),
            ]
            
            for method_name, method_func in conversion_methods:
                try:
                    logger.info(f"Trying conversion with {method_name}")
                    conversion_success = method_func(input_path, output_path)
                    if conversion_success:
                        logger.info(f"Conversion with {method_name} successful")
                        break
                except Exception as e:
                    error_msg = f"{method_name} conversion failed: {str(e)}"
                    logger.warning(error_msg)
                    error_messages.append(error_msg)
            
            # If no conversion method succeeded
            if not conversion_success:
                logger.error("All conversion methods failed")
                raise Exception("All conversion methods failed: " + " | ".join(error_messages))
            
            # Clean up the uploaded file
            os.remove(input_path)
            
            # Return the download URL
            return jsonify({
                'success': True,
                'filename': os.path.splitext(file.filename)[0] + '.pdf',
                'download_url': f'/download/{output_filename}'
            })
            
        except Exception as e:
            logger.error(f"Conversion error: {str(e)}")
            # If conversion fails, clean up and return error
            if os.path.exists(input_path):
                os.remove(input_path)
            return jsonify({'error': str(e)}), 500
    else:
        return jsonify({'error': 'Invalid file format. Please upload a Word document (.doc or .docx)'}), 400

def convert_with_pandoc(input_path, output_path):
    """Convert a Word document to PDF using Pandoc (high quality)"""
    try:
        # Check if pandoc is installed
        if not shutil.which('pandoc'):
            raise FileNotFoundError("Pandoc not found. Install with 'apt-get install pandoc texlive-latex-base texlive-fonts-recommended'")
        
        # Run the conversion with Pandoc (via LaTeX for best results)
        cmd = ['pandoc', input_path, '--pdf-engine=xelatex', '-o', output_path]
        process = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        
        if process.returncode != 0:
            raise Exception(f"Pandoc returned error code {process.returncode}: {process.stderr}")
        
        return os.path.exists(output_path) and os.path.getsize(output_path) > 0
        
    except Exception as e:
        raise Exception(f"Pandoc conversion error: {str(e)}")

def find_libreoffice_executable():
    """Find the LibreOffice executable on the system"""
    possible_executables = [
        'libreoffice',
        'soffice',
        '/usr/bin/libreoffice',
        '/usr/bin/soffice',
        '/usr/lib/libreoffice/program/soffice',
        '/opt/libreoffice*/program/soffice',
        '/Applications/LibreOffice.app/Contents/MacOS/soffice'  # macOS
    ]
    
    for exe in possible_executables:
        # Handle wildcard paths
        if '*' in exe:
            import glob
            matches = glob.glob(exe)
            for match in matches:
                if os.path.isfile(match) and os.access(match, os.X_OK):
                    return match
        # Direct path or command name
        elif exe and shutil.which(exe):
            return shutil.which(exe)
    
    raise FileNotFoundError("LibreOffice executable not found. Please install LibreOffice.")

def convert_with_libreoffice(input_path, output_path):
    """Convert a Word document to PDF using LibreOffice"""
    output_dir = os.path.dirname(output_path)
    
    try:
        # Find LibreOffice executable
        libreoffice_exec = find_libreoffice_executable()
        
        # Command to convert using LibreOffice with better parameters
        cmd = [
            libreoffice_exec, 
            '--headless', 
            '--norestore',
            '--invisible',
            '--convert-to', 
            'pdf:writer_pdf_Export',  # Use the PDF/A-1a format for better quality
            '--outdir', 
            output_dir,
            input_path
        ]
        
        process = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        
        if process.returncode != 0:
            raise Exception(f"LibreOffice returned error code {process.returncode}: {process.stderr}")
        
        # LibreOffice puts the output in the output_dir with original name but .pdf extension
        base_filename = os.path.basename(input_path)
        base_name_no_ext = os.path.splitext(base_filename)[0]
        temp_output = os.path.join(output_dir, f"{base_name_no_ext}.pdf")
        
        if os.path.exists(temp_output) and temp_output != output_path:
            os.rename(temp_output, output_path)
            return os.path.getsize(output_path) > 0
        elif os.path.exists(output_path):
            return os.path.getsize(output_path) > 0
        else:
            raise FileNotFoundError("Conversion completed but output file not found")
            
    except FileNotFoundError as e:
        raise e
    except Exception as e:
        raise Exception(f"LibreOffice conversion error: {str(e)}")

def convert_with_unoconv(input_path, output_path):
    """Convert a Word document to PDF using unoconv (another LibreOffice-based tool)"""
    try:
        # Check if unoconv is installed
        if not shutil.which('unoconv'):
            raise FileNotFoundError("unoconv not found. Install with 'pip install unoconv' or 'apt-get install unoconv'")
        
        # Run the conversion with better parameters
        cmd = ['unoconv', '-f', 'pdf', '--format=pdf', '-eSelectPdfVersion=1', '-o', output_path, input_path]
        process = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        
        if process.returncode != 0:
            raise Exception(f"unoconv returned error code {process.returncode}: {process.stderr}")
        
        return os.path.exists(output_path) and os.path.getsize(output_path) > 0
        
    except Exception as e:
        raise Exception(f"unoconv conversion error: {str(e)}")

def convert_with_python_docx(input_path, output_path):
    """
    Convert using python-docx for reading and PyPDF for writing.
    Better handling of document structure compared to the reportlab version.
    """
    try:
        # Import necessary libraries
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet
        
        # Open the Word document
        doc = Document(input_path)
        
        # Create PDF document
        pdf = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Extract content from Word document
        content = []
        
        # Process paragraphs
        for para in doc.paragraphs:
            if not para.text.strip():
                content.append(Spacer(1, 12))
                continue
            
            # Determine style based on paragraph properties
            style = 'Normal'
            if para.style.name.startswith('Heading'):
                style = 'Heading1'
            
            # Add paragraph to content
            content.append(Paragraph(para.text, styles[style]))
            content.append(Spacer(1, 6))
            
        # Process tables (simplified handling)
        for table in doc.tables:
            data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                data.append(row_data)
            
            if data:  # Only process if there's actually data
                table_style = TableStyle([
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ])
                content.append(Table(data, style=table_style))
                content.append(Spacer(1, 12))
        
        # Build PDF
        pdf.build(content)
        
        return os.path.exists(output_path) and os.path.getsize(output_path) > 0
        
    except ImportError:
        raise Exception("Required libraries not installed. Run 'pip install python-docx reportlab'")
    except Exception as e:
        raise Exception(f"python-docx conversion error: {str(e)}")

@app.route('/download/<filename>')
def download_file(filename):
    file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    if not os.path.exists(file_path):
        return jsonify({'error': 'File not found'}), 404
        
    return send_file(file_path, 
                    as_attachment=True, 
                    download_name=filename.split('_', 2)[2])  # Remove the unique ID prefix

# Clean up old files periodically
@app.route('/cleanup', methods=['POST'])
def cleanup():
    threshold_time = time.time() - (24 * 60 * 60)  # 24 hours ago
    
    deleted_count = 0
    for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
        for filename in os.listdir(folder):
            file_path = os.path.join(folder, filename)
            try:
                # Extract timestamp from filename
                timestamp = int(filename.split('_')[0])
                if timestamp < threshold_time:
                    os.remove(file_path)
                    deleted_count += 1
            except (ValueError, IndexError, OSError):
                # Skip files that don't follow the naming convention or can't be deleted
                continue
    
    return jsonify({'success': True, 'deleted_count': deleted_count})

if __name__ == '__main__':
    app.run(debug=True)
